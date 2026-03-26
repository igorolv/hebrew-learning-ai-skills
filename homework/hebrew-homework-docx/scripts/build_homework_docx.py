#!/usr/bin/env python3
"""Build Hebrew homework markdown into formatted DOCX."""
from __future__ import annotations

import argparse
import math
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# ---------------------------------------------------------------------------
# Constants / regexes
# ---------------------------------------------------------------------------

HEBREW_RE = re.compile(r"[\u0590-\u05FF\uFB1D-\uFB4F]")
HEADING1_RE = re.compile(r"^#\s+(.+)$")
HEADING2_RE = re.compile(r"^##\s+(.+)$")
HEADING3_RE = re.compile(r"^###\s+(.+)$")
SEPARATOR_RE = re.compile(r"^---\s*$")
TABLE_ROW_RE = re.compile(r"^\|.*\|\s*$")
TABLE_DELIM_RE = re.compile(r"^\|\s*:?[-]+:?(?:\s*\|\s*:?[-]+:?)+\s*\|\s*$")
NUMBERED_RE = re.compile(r"^(\d+)\.\s+(.+)$")
EXERCISE_HEADER_RE = re.compile(r"^\|\s*Текст\s*\|\s*№\s*\|")
INLINE_RE = re.compile(r"(\*\*(.+?)\*\*)|(\*(.+?)\*)|([^*]+)")

# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------


@dataclass
class Run:
    text: str
    bold: bool = False
    italic: bool = False
    hebrew: bool = False


# ---------------------------------------------------------------------------
# Language helpers
# ---------------------------------------------------------------------------


def _is_hebrew(ch: str) -> bool:
    c = ord(ch)
    return (0x0590 <= c <= 0x05FF) or (0xFB1D <= c <= 0xFB4F)


def _is_cyrillic(ch: str) -> bool:
    c = ord(ch)
    return 0x0400 <= c <= 0x04FF


def contains_hebrew(text: str) -> bool:
    return bool(HEBREW_RE.search(text or ""))


def is_all_hebrew(text: str) -> bool:
    """True when every letter-like character is Hebrew."""
    has_heb = False
    for ch in text:
        if ch.isspace() or not ch.isalpha() and ch not in "\u05B0\u05B1\u05B2\u05B3\u05B4\u05B5\u05B6\u05B7\u05B8\u05B9\u05BA\u05BB\u05BC\u05BD\u05BF\u05C1\u05C2\u05C4\u05C5\u05C7":
            if _is_hebrew(ch):
                has_heb = True
            continue
        if _is_hebrew(ch):
            has_heb = True
        elif _is_cyrillic(ch) or (ch.isascii() and ch.isalpha()):
            return False
    return has_heb


def paragraph_direction(text: str) -> str:
    """Return 'rtl' or 'ltr' based on first strong directional character."""
    for ch in text:
        if _is_hebrew(ch):
            return "rtl"
        if _is_cyrillic(ch) or (ch.isascii() and ch.isalpha()):
            return "ltr"
    return "ltr"


# ---------------------------------------------------------------------------
# Inline formatting: **bold**, *italic*, script splitting
# ---------------------------------------------------------------------------


def split_by_script(text: str) -> List[tuple[str, bool]]:
    """Split *text* into segments of same script → [(text, is_hebrew), ...]."""
    if not text:
        return []
    segments: list[tuple[str, bool]] = []
    buf = ""
    cur: Optional[bool] = None

    for ch in text:
        if _is_hebrew(ch):
            ch_heb = True
        elif _is_cyrillic(ch) or (ch.isascii() and ch.isalpha()):
            ch_heb = False
        else:
            buf += ch
            continue

        if cur is None:
            cur = ch_heb
        if ch_heb != cur:
            if buf:
                segments.append((buf, cur))
            buf = ch
            cur = ch_heb
        else:
            buf += ch

    if buf:
        segments.append((buf, cur if cur is not None else False))
    return segments


def parse_inline(text: str) -> List[Run]:
    """Parse **bold** / *italic* then split each piece by script."""
    raw: list[Run] = []
    for m in INLINE_RE.finditer(text):
        if m.group(2):
            raw.append(Run(text=m.group(2), bold=True))
        elif m.group(4):
            raw.append(Run(text=m.group(4), italic=True))
        elif m.group(5):
            raw.append(Run(text=m.group(5)))

    result: list[Run] = []
    for r in raw:
        for seg_text, seg_heb in split_by_script(r.text):
            result.append(Run(text=seg_text, bold=r.bold, italic=r.italic, hebrew=seg_heb))
    return result


# ---------------------------------------------------------------------------
# Line classification
# ---------------------------------------------------------------------------


def classify_lines(lines: List[str]) -> List[tuple[str, str]]:
    """Return [(kind, content), ...] for each line."""
    result: list[tuple[str, str]] = []
    for line in lines:
        s = line.strip()
        if not s:
            result.append(("blank", ""))
            continue
        if SEPARATOR_RE.match(s):
            result.append(("separator", ""))
            continue
        m = HEADING1_RE.match(s)
        if m:
            result.append(("heading1", m.group(1)))
            continue
        m = HEADING2_RE.match(s)
        if m:
            result.append(("heading2", m.group(1)))
            continue
        m = HEADING3_RE.match(s)
        if m:
            result.append(("heading3", m.group(1)))
            continue
        if TABLE_ROW_RE.match(s):
            result.append(("table_row", s))
            continue
        m = NUMBERED_RE.match(s)
        if m:
            result.append(("numbered", s))
            continue
        if is_all_hebrew(s):
            result.append(("hebrew_paragraph", s))
            continue
        result.append(("text_paragraph", s))
    return result


# ---------------------------------------------------------------------------
# Table helpers
# ---------------------------------------------------------------------------


def normalize_table_row(line: str) -> List[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_exercise_table(header_line: str) -> bool:
    return bool(EXERCISE_HEADER_RE.match(header_line.strip()))


# ---------------------------------------------------------------------------
# OOXML low-level helpers (mirrors hp-generate-docx patterns)
# ---------------------------------------------------------------------------


def _set_bidi(paragraph, enabled: bool) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    el = pPr.find(qn("w:bidi"))
    if el is None:
        el = OxmlElement("w:bidi")
        pPr.append(el)
    el.set(qn("w:val"), "1" if enabled else "0")


def _set_jc(paragraph, value: str) -> None:
    """Set w:jc directly: 'start', 'end', 'center', 'both', 'left', 'right'."""
    pPr = paragraph._p.get_or_add_pPr()
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), value)


def _set_rtl_run(run, rtl: bool) -> None:
    rPr = run._r.get_or_add_rPr()
    el = rPr.find(qn("w:rtl"))
    if el is None:
        el = OxmlElement("w:rtl")
        rPr.append(el)
    el.set(qn("w:val"), "1" if rtl else "0")


def _set_run_font(run, *, font: str, size_pt: int, rtl: bool) -> None:
    run.font.name = font
    run.font.size = Pt(size_pt)
    rPr = run._r.get_or_add_rPr()

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), font)

    half = str(size_pt * 2)
    for tag in ("w:sz", "w:szCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:val"), half)

    _set_rtl_run(run, rtl)


def _set_bold_cs(run) -> None:
    rPr = run._r.get_or_add_rPr()
    if rPr.find(qn("w:bCs")) is None:
        rPr.append(OxmlElement("w:bCs"))


def _set_italic_cs(run) -> None:
    rPr = run._r.get_or_add_rPr()
    if rPr.find(qn("w:iCs")) is None:
        rPr.append(OxmlElement("w:iCs"))


def _set_cell_width(cell, twips: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    w = tcPr.find(qn("w:tcW"))
    if w is None:
        w = OxmlElement("w:tcW")
        tcPr.append(w)
    w.set(qn("w:w"), str(twips))
    w.set(qn("w:type"), "dxa")


def _set_row_cant_split(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def _text_width_twips(doc: Document) -> int:
    sec = doc.sections[-1]
    return int(round((sec.page_width - sec.left_margin - sec.right_margin) / 635))


# ---------------------------------------------------------------------------
# Run builder: applies font/style to a docx Run based on our Run dataclass
# ---------------------------------------------------------------------------


def _apply_run_style(docx_run, r: Run) -> None:
    if r.hebrew:
        _set_run_font(docx_run, font="David", size_pt=18, rtl=True)
        if r.bold:
            docx_run.bold = True
            _set_bold_cs(docx_run)
        if r.italic:
            docx_run.italic = True
            _set_italic_cs(docx_run)
    else:
        _set_run_font(docx_run, font="Arial", size_pt=10, rtl=False)
        if r.bold:
            docx_run.bold = True
        if r.italic:
            docx_run.italic = True


def _add_runs(paragraph, runs: List[Run]) -> None:
    for r in runs:
        _apply_run_style(paragraph.add_run(r.text), r)


def _align_paragraph(paragraph, text: str) -> None:
    """Set alignment + bidi based on first strong character."""
    if paragraph_direction(text) == "rtl":
        _set_bidi(paragraph, True)
        _set_jc(paragraph, "start")
    else:
        _set_bidi(paragraph, False)
        _set_jc(paragraph, "start")


# ---------------------------------------------------------------------------
# Paragraph builders
# ---------------------------------------------------------------------------


def add_heading1(doc: Document, text: str) -> None:
    """# Домашнее задание — Урок N, часть M"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.bold = True
    _set_run_font(r, font="Arial", size_pt=14, rtl=False)


def add_heading2(doc: Document, text: str, first: bool) -> None:
    """## Слайд X — [тип]  (page break before, except first)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    if not first:
        p.paragraph_format.page_break_before = True
    r = p.add_run(text)
    r.bold = True
    _set_run_font(r, font="Arial", size_pt=13, rtl=False)


def add_heading3(doc: Document, text: str) -> None:
    """### subheading — Hebrew (David 16pt bold RTL) or Russian (Arial 10pt bold)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if contains_hebrew(text):
        _set_bidi(p, True)
        _set_jc(p, "start")
        r = p.add_run(text)
        r.bold = True
        _set_run_font(r, font="David", size_pt=16, rtl=True)
        _set_bold_cs(r)
    else:
        _set_jc(p, "start")
        r = p.add_run(text)
        r.bold = True
        _set_run_font(r, font="Arial", size_pt=10, rtl=False)


def add_hebrew_paragraph(doc: Document, text: str) -> None:
    """Standalone Hebrew text (instruction) — David 18pt bold RTL."""
    p = doc.add_paragraph()
    _set_bidi(p, True)
    _set_jc(p, "start")
    p.paragraph_format.space_after = Pt(6)
    runs = parse_inline(text)
    for r in runs:
        dr = p.add_run(r.text)
        _set_run_font(dr, font="David", size_pt=18, rtl=True)
        dr.bold = True
        _set_bold_cs(dr)


def add_text_paragraph(doc: Document, text: str) -> None:
    """Regular / mixed paragraph."""
    p = doc.add_paragraph()
    _align_paragraph(p, text)
    p.paragraph_format.space_after = Pt(6)
    _add_runs(p, parse_inline(text))


def add_numbered_item(doc: Document, text: str) -> None:
    """Legacy format: 1. [text]  — number + content."""
    m = NUMBERED_RE.match(text.strip())
    if not m:
        add_text_paragraph(doc, text)
        return
    num, body = m.group(1), m.group(2)
    p = doc.add_paragraph()
    _align_paragraph(p, body)
    p.paragraph_format.space_after = Pt(6)
    # number run
    nr = p.add_run(f"{num}. ")
    _set_run_font(nr, font="Arial", size_pt=10, rtl=False)
    # body runs
    _add_runs(p, parse_inline(body))


# ---------------------------------------------------------------------------
# Table builder
# ---------------------------------------------------------------------------


def _style_cell(cell, text: str, force_hebrew: bool = False) -> None:
    """Format a table cell with parsed inline runs."""
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if force_hebrew:
        _set_bidi(p, True)
        _set_jc(p, "start")
        if text.strip():
            r = p.add_run(text)
            _set_run_font(r, font="David", size_pt=18, rtl=True)
    else:
        _set_bidi(p, True)
        _set_jc(p, "start")
        _add_runs(p, parse_inline(text))


def add_table(doc: Document, rows: List[List[str]], exercise: bool) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    # fixed layout
    tPr = tbl._tbl.tblPr
    lay = tPr.find(qn("w:tblLayout"))
    if lay is None:
        lay = OxmlElement("w:tblLayout")
        tPr.append(lay)
    lay.set(qn("w:type"), "fixed")

    tw = _text_width_twips(doc)
    if exercise and ncols == 2:
        widths = [int(tw * 0.95), int(tw * 0.05)]
    else:
        cw = max(1, math.floor(tw / ncols))
        widths = [cw] * ncols

    for ri, rd in enumerate(rows):
        row = tbl.rows[ri]
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        _set_row_cant_split(row)
        for ci in range(ncols):
            cell = row.cells[ci]
            _set_cell_width(cell, widths[ci])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # remove noWrap if any
            tcPr = cell._tc.get_or_add_tcPr()
            nw = tcPr.find(qn("w:noWrap"))
            if nw is not None:
                tcPr.remove(nw)
            txt = rd[ci] if ci < len(rd) else ""
            is_num_col = exercise and ncols == 2 and ci == 1
            _style_cell(cell, txt, force_hebrew=is_num_col)
            # cell margins
            mar = tcPr.find(qn("w:tcMar"))
            if mar is None:
                mar = OxmlElement("w:tcMar")
                tcPr.append(mar)
            for side in ("top", "left", "bottom", "right"):
                el = mar.find(qn(f"w:{side}"))
                if el is None:
                    el = OxmlElement(f"w:{side}")
                    mar.append(el)
                el.set(qn("w:w"), "100")
                el.set(qn("w:type"), "dxa")


# ---------------------------------------------------------------------------
# Document defaults
# ---------------------------------------------------------------------------


def _init_document(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.27)
    sec.bottom_margin = Cm(1.27)
    sec.left_margin = Cm(1.27)
    sec.right_margin = Cm(1.27)


# ---------------------------------------------------------------------------
# Validation
# ---------------------------------------------------------------------------


def validate(text: str) -> List[str]:
    errors: list[str] = []
    if not text.strip().startswith("# Домашнее задание"):
        errors.append("Файл должен начинаться с '# Домашнее задание'")
    if not re.search(r"^##\s+Слайд", text, re.MULTILINE):
        errors.append("Файл должен содержать хотя бы одну секцию '## Слайд'")
    return errors


# ---------------------------------------------------------------------------
# Main build
# ---------------------------------------------------------------------------


def build(md_path: Path, output: Optional[Path] = None) -> Path:
    text = md_path.read_text(encoding="utf-8")
    errors = validate(text)
    if errors:
        for e in errors:
            print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)

    raw_lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    classified = classify_lines(raw_lines)

    doc = Document()
    _init_document(doc)

    h2_count = 0
    i = 0
    n = len(classified)

    while i < n:
        kind, content = classified[i]

        if kind == "heading1":
            add_heading1(doc, content)
            i += 1
            continue

        if kind == "heading2":
            h2_count += 1
            add_heading2(doc, content, first=(h2_count == 1))
            i += 1
            continue

        if kind == "heading3":
            add_heading3(doc, content)
            i += 1
            continue

        if kind in ("separator", "blank"):
            i += 1
            continue

        if kind == "table_row":
            tbl_lines: list[str] = []
            while i < n and classified[i][0] == "table_row":
                tbl_lines.append(classified[i][1])
                i += 1
            exercise = is_exercise_table(tbl_lines[0]) if tbl_lines else False
            rows = [normalize_table_row(ln) for ln in tbl_lines]
            has_header = len(rows) >= 2 and TABLE_DELIM_RE.match(tbl_lines[1])
            if has_header:
                all_rows = rows[2:]  # drop header + delimiter
            else:
                all_rows = rows
            if all_rows:
                add_table(doc, all_rows, exercise)
            continue

        if kind == "hebrew_paragraph":
            add_hebrew_paragraph(doc, content)
            i += 1
            continue

        if kind == "numbered":
            add_numbered_item(doc, content)
            i += 1
            continue

        # text_paragraph
        add_text_paragraph(doc, content)
        i += 1

    out = output or md_path.with_suffix(".docx")
    doc.save(str(out))
    return out


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def parse_args(argv: Optional[Iterable[str]] = None) -> argparse.Namespace:
    ap = argparse.ArgumentParser(description="Build Hebrew homework markdown into DOCX.")
    ap.add_argument("markdown", type=Path, nargs="+", help="Homework .md file(s)")
    ap.add_argument("-o", "--output", type=Path, help="Output DOCX (single input only)")
    return ap.parse_args(argv)


def main(argv: Optional[Iterable[str]] = None) -> int:
    args = parse_args(argv)
    if args.output and len(args.markdown) > 1:
        print("ERROR: -o/--output only works with a single input", file=sys.stderr)
        return 1
    for md in args.markdown:
        if not md.exists():
            print(f"ERROR: not found: {md}", file=sys.stderr)
            return 1
        out = build(md, args.output)
        print(f"DOCX: {out}", flush=True, file=open(sys.stdout.fileno(), "w", encoding="utf-8", closefd=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
