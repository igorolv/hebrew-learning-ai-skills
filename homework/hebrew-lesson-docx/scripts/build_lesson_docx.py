#!/usr/bin/env python3
"""Build Hebrew lesson cheat-sheet markdown into formatted DOCX."""
from __future__ import annotations

import argparse
import math
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


HEBREW_RE = re.compile(r"[\u0590-\u05FF\uFB1D-\uFB4F]")
DOC_TITLE_RE = re.compile(r"^#\s+(Шпаргалка|Грамматика)\b(.+)?$")
DOC_SUBTITLE_RE = re.compile(r"^##\s+שיעור\b(.+)?$")
SECTION_H1_RE = re.compile(r"^#\s+(\d+\.\s+.+)$")
HEADING2_RE = re.compile(r"^##\s+(.+)$")
HEADING3_RE = re.compile(r"^###\s+(.+)$")
BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)$")
BULLET_RE = re.compile(r"^([*-])\s+(.+)$")
SEPARATOR_RE = re.compile(r"^---\s*$")
TABLE_ROW_RE = re.compile(r"^\|.*\|\s*$")
TABLE_DELIM_RE = re.compile(r"^\|\s*:?[-]+:?(?:\s*\|\s*:?[-]+:?)+\s*\|\s*$")
INLINE_RE = re.compile(r"(\*\*(.+?)\*\*)|(\*(.+?)\*)|([^*]+)")


@dataclass
class Run:
    text: str
    bold: bool = False
    italic: bool = False
    hebrew: bool = False


def _is_hebrew(ch: str) -> bool:
    c = ord(ch)
    return (0x0590 <= c <= 0x05FF) or (0xFB1D <= c <= 0xFB4F)


def _is_cyrillic(ch: str) -> bool:
    c = ord(ch)
    return 0x0400 <= c <= 0x04FF


def contains_hebrew(text: str) -> bool:
    return bool(HEBREW_RE.search(text or ""))


def is_all_hebrew(text: str) -> bool:
    has_heb = False
    for ch in text:
        if ch.isspace() or (not ch.isalpha() and not _is_hebrew(ch)):
            if _is_hebrew(ch):
                has_heb = True
            continue
        if _is_hebrew(ch):
            has_heb = True
        elif _is_cyrillic(ch) or (ch.isascii() and ch.isalpha()):
            return False
    return has_heb


def count_scripts(text: str) -> tuple[int, int]:
    heb = 0
    other = 0
    for ch in text:
        if _is_hebrew(ch):
            heb += 1
        elif _is_cyrillic(ch) or (ch.isascii() and ch.isalpha()):
            other += 1
    return heb, other


def dominant_direction(text: str) -> str:
    heb, other = count_scripts(text)
    return "rtl" if heb > other else "ltr"


def split_by_script(text: str) -> List[tuple[str, bool]]:
    if not text:
        return []

    segments: list[tuple[str, bool]] = []
    buf = ""
    cur: Optional[bool] = None
    neutral = ""

    for ch in text:
        if _is_hebrew(ch):
            script = True
        elif _is_cyrillic(ch) or (ch.isascii() and ch.isalpha()):
            script = False
        else:
            neutral += ch
            continue

        if cur is None:
            if neutral:
                segments.append((neutral, False))
                neutral = ""
            cur = script
            buf = ch
            continue

        if script == cur:
            if neutral:
                buf += neutral
                neutral = ""
            buf += ch
            continue

        if neutral:
            if cur:
                if buf:
                    segments.append((buf, cur))
                buf = neutral + ch
            else:
                buf += neutral
                if buf:
                    segments.append((buf, cur))
                buf = ch
            neutral = ""
        else:
            if buf:
                segments.append((buf, cur))
            buf = ch
        cur = script

    if cur is None:
        if neutral:
            segments.append((neutral, False))
    else:
        if neutral:
            buf += neutral
        if buf:
            segments.append((buf, cur))

    return segments


def parse_inline(text: str) -> List[Run]:
    raw: list[Run] = []
    for match in INLINE_RE.finditer(text):
        if match.group(2):
            raw.append(Run(text=match.group(2), bold=True))
        elif match.group(4):
            raw.append(Run(text=match.group(4), italic=True))
        elif match.group(5):
            raw.append(Run(text=match.group(5)))

    result: list[Run] = []
    for run in raw:
        for seg_text, seg_heb in split_by_script(run.text):
            result.append(Run(text=seg_text, bold=run.bold, italic=run.italic, hebrew=seg_heb))

    for idx in range(len(result) - 1):
        current = result[idx]
        following = result[idx + 1]
        if not current.hebrew or following.hebrew:
            continue
        match = re.search(r"(\s+)$", current.text)
        if not match:
            continue
        spacer = match.group(1)
        current.text = current.text[: -len(spacer)]
        following.text = spacer + following.text

    for idx in range(len(result) - 1):
        current = result[idx]
        following = result[idx + 1]
        if not current.hebrew or following.hebrew:
            continue
        match = re.match(r"^(\s+)(.*)$", following.text, re.DOTALL)
        if not match:
            continue
        spacer = match.group(1).replace(" ", "\u00A0")
        following.text = spacer + match.group(2)

    return [run for run in result if run.text]


def classify_lines(lines: List[str]) -> List[tuple[str, str]]:
    result: list[tuple[str, str]] = []
    for line in lines:
        s = line.rstrip()
        stripped = s.strip()
        if not stripped:
            result.append(("blank", ""))
            continue
        if SEPARATOR_RE.match(stripped):
            result.append(("separator", ""))
            continue
        if DOC_TITLE_RE.match(stripped):
            result.append(("doc_title", stripped[2:].strip()))
            continue
        if DOC_SUBTITLE_RE.match(stripped):
            result.append(("doc_subtitle", stripped[3:].strip()))
            continue
        match = SECTION_H1_RE.match(stripped)
        if match:
            result.append(("section_heading", match.group(1)))
            continue
        match = HEADING3_RE.match(stripped)
        if match:
            result.append(("heading3", match.group(1)))
            continue
        match = HEADING2_RE.match(stripped)
        if match:
            result.append(("heading2", match.group(1)))
            continue
        match = BLOCKQUOTE_RE.match(stripped)
        if match:
            result.append(("blockquote", match.group(1)))
            continue
        match = BULLET_RE.match(stripped)
        if match:
            result.append(("bullet", match.group(2)))
            continue
        if TABLE_ROW_RE.match(stripped):
            result.append(("table_row", stripped))
            continue
        if is_all_hebrew(stripped):
            result.append(("hebrew_paragraph", stripped))
            continue
        result.append(("text_paragraph", stripped))
    return result


def normalize_table_row(line: str) -> List[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _set_bidi(paragraph, enabled: bool) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    el = p_pr.find(qn("w:bidi"))
    if el is None:
        el = OxmlElement("w:bidi")
        p_pr.append(el)
    el.set(qn("w:val"), "1" if enabled else "0")


def _set_jc(paragraph, value: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), value)


def _set_rtl_run(run, rtl: bool) -> None:
    r_pr = run._r.get_or_add_rPr()
    el = r_pr.find(qn("w:rtl"))
    if el is None:
        el = OxmlElement("w:rtl")
        r_pr.append(el)
    el.set(qn("w:val"), "1" if rtl else "0")


def _set_cs_run(run, enabled: bool) -> None:
    r_pr = run._r.get_or_add_rPr()
    el = r_pr.find(qn("w:cs"))
    if enabled:
        if el is None:
            el = OxmlElement("w:cs")
            r_pr.append(el)
        el.set(qn("w:val"), "1")
    elif el is not None:
        r_pr.remove(el)


def _set_run_font(run, *, font: str, size_pt: int, rtl: bool) -> None:
    run.font.name = font
    run.font.size = Pt(size_pt)
    r_pr = run._r.get_or_add_rPr()

    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), font)

    half_points = str(size_pt * 2)
    for tag in ("w:sz", "w:szCs"):
        el = r_pr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            r_pr.append(el)
        el.set(qn("w:val"), half_points)

    _set_rtl_run(run, rtl)
    _set_cs_run(run, rtl)


def _set_bold_cs(run) -> None:
    r_pr = run._r.get_or_add_rPr()
    if r_pr.find(qn("w:bCs")) is None:
        r_pr.append(OxmlElement("w:bCs"))


def _set_italic_cs(run) -> None:
    r_pr = run._r.get_or_add_rPr()
    if r_pr.find(qn("w:iCs")) is None:
        r_pr.append(OxmlElement("w:iCs"))


def _set_paragraph_left_border(paragraph, color: str = "4472C4", size: str = "12", space: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)


def _set_cell_width(cell, twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    w = tc_pr.find(qn("w:tcW"))
    if w is None:
        w = OxmlElement("w:tcW")
        tc_pr.append(w)
    w.set(qn("w:w"), str(twips))
    w.set(qn("w:type"), "dxa")


def _set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")


def _shade_cell(cell, fill: str = "D9E2F3") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _text_width_twips(doc: Document) -> int:
    sec = doc.sections[-1]
    return int(round((sec.page_width - sec.left_margin - sec.right_margin) / 635))


def _apply_run_style(docx_run, run: Run, *, hebrew_pt: int, other_pt: int, force_bold: bool = False, force_italic: bool = False) -> None:
    bold = force_bold or run.bold
    italic = force_italic or run.italic

    if run.hebrew:
        _set_run_font(docx_run, font="David", size_pt=hebrew_pt, rtl=True)
        if bold:
            docx_run.bold = True
            _set_bold_cs(docx_run)
        if italic:
            docx_run.italic = True
            _set_italic_cs(docx_run)
    else:
        _set_run_font(docx_run, font="Arial", size_pt=other_pt, rtl=False)
        if bold:
            docx_run.bold = True
        if italic:
            docx_run.italic = True


def _add_runs(paragraph, runs: List[Run], *, hebrew_pt: int, other_pt: int, force_bold: bool = False, force_italic: bool = False) -> None:
    for run in runs:
        _apply_run_style(
            paragraph.add_run(run.text),
            run,
            hebrew_pt=hebrew_pt,
            other_pt=other_pt,
            force_bold=force_bold,
            force_italic=force_italic,
        )


def _align_paragraph(paragraph, text: str) -> None:
    if dominant_direction(text) == "rtl":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_bidi(paragraph, True)
        _set_jc(paragraph, "right")
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_bidi(paragraph, False)
        _set_jc(paragraph, "left")


def add_doc_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True
    _set_run_font(run, font="Arial", size_pt=16, rtl=False)


def add_doc_subtitle(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(10)
    _set_bidi(paragraph, True)
    _set_jc(paragraph, "center")
    run = paragraph.add_run(text)
    run.bold = True
    _set_run_font(run, font="David", size_pt=16, rtl=True)
    _set_bold_cs(run)


def add_section_heading(doc: Document, text: str, *, page_break_before: bool) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.page_break_before = page_break_before
    _set_bidi(paragraph, False)
    _set_jc(paragraph, "left")
    _add_runs(paragraph, parse_inline(text), hebrew_pt=14, other_pt=14, force_bold=True)


def add_heading2(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(5)
    if dominant_direction(text) == "rtl":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_bidi(paragraph, True)
        _set_jc(paragraph, "right")
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_bidi(paragraph, False)
        _set_jc(paragraph, "left")
    _add_runs(paragraph, parse_inline(text), hebrew_pt=13, other_pt=13, force_bold=True)


def add_heading3(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    _align_paragraph(paragraph, text)
    _add_runs(paragraph, parse_inline(text), hebrew_pt=12, other_pt=12, force_bold=True)


def add_text_paragraph(doc: Document, text: str, *, space_after: int = 80) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after / 20)
    _align_paragraph(paragraph, text)
    _add_runs(paragraph, parse_inline(text), hebrew_pt=18, other_pt=12)


def add_hebrew_paragraph(doc: Document, text: str, *, space_after: int = 80) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after / 20)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_bidi(paragraph, True)
    _set_jc(paragraph, "right")
    _add_runs(paragraph, parse_inline(text), hebrew_pt=18, other_pt=12)


def add_contrast_pair(doc: Document, hebrew: str, russian: str) -> None:
    add_hebrew_paragraph(doc, hebrew, space_after=0)
    add_text_paragraph(doc, russian, space_after=120)


def add_blockquote(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(1.27)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    _set_paragraph_left_border(paragraph)
    _align_paragraph(paragraph, text)
    _add_runs(paragraph, parse_inline(text), hebrew_pt=18, other_pt=12, force_bold=True)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(1.27)
    paragraph.paragraph_format.first_line_indent = Cm(-0.635)
    paragraph.paragraph_format.space_after = Pt(2)
    bullet = paragraph.add_run("• ")
    _set_run_font(bullet, font="Arial", size_pt=12, rtl=False)
    _add_runs(paragraph, parse_inline(text), hebrew_pt=18, other_pt=12)


def _style_cell(cell, text: str, *, is_header: bool) -> None:
    for paragraph in cell.paragraphs:
        paragraph.clear()

    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    _align_paragraph(paragraph, text)

    runs = parse_inline(text)
    if runs:
        _add_runs(
            paragraph,
            runs,
            hebrew_pt=16 if is_header else 18,
            other_pt=12,
            force_bold=is_header,
        )
    else:
        paragraph.add_run("")


def add_table(doc: Document, rows: List[List[str]]) -> None:
    if not rows:
        return

    ncols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    _set_table_borders(table)

    text_width = _text_width_twips(doc)
    col_width = max(1, math.floor(text_width / ncols))
    widths = [col_width] * ncols

    for row_index, row_data in enumerate(rows):
        row = table.rows[row_index]
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        _set_row_cant_split(row)

        for col_index in range(ncols):
            cell = row.cells[col_index]
            _set_cell_width(cell, widths[col_index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            mar = tc_pr.find(qn("w:tcMar"))
            if mar is None:
                mar = OxmlElement("w:tcMar")
                tc_pr.append(mar)
            for side in ("top", "left", "bottom", "right"):
                el = mar.find(qn(f"w:{side}"))
                if el is None:
                    el = OxmlElement(f"w:{side}")
                    mar.append(el)
                el.set(qn("w:w"), "60")
                el.set(qn("w:type"), "dxa")

            text = row_data[col_index] if col_index < len(row_data) else ""
            _style_cell(cell, text, is_header=(row_index == 0))
            if row_index == 0:
                _shade_cell(cell)


def _init_document(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


def validate(text: str) -> List[str]:
    errors: list[str] = []
    stripped = text.lstrip()
    if not (stripped.startswith("# Шпаргалка") or stripped.startswith("# Грамматика")):
        errors.append("Файл должен начинаться с '# Шпаргалка' или '# Грамматика'")
    if not re.search(r"^#\s+\d+\.\s+", text, re.MULTILINE):
        errors.append("Файл должен содержать хотя бы один нумерованный раздел '# N.'")
    return errors


def build(md_path: Path, output: Optional[Path] = None) -> Path:
    text = md_path.read_text(encoding="utf-8")
    errors = validate(text)
    if errors:
        for error in errors:
            print(f"ERROR: {error}", file=sys.stderr)
        sys.exit(1)

    raw_lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    classified = classify_lines(raw_lines)

    doc = Document()
    _init_document(doc)

    section_count = 0
    index = 0
    total = len(classified)

    while index < total:
        kind, content = classified[index]

        if kind in ("blank", "separator"):
            index += 1
            continue

        if kind == "doc_title":
            add_doc_title(doc, content)
            index += 1
            continue

        if kind == "doc_subtitle":
            add_doc_subtitle(doc, content)
            index += 1
            continue

        if kind == "section_heading":
            section_count += 1
            add_section_heading(doc, content, page_break_before=(section_count > 1))
            index += 1
            continue

        if kind == "heading2":
            add_heading2(doc, content)
            index += 1
            continue

        if kind == "heading3":
            add_heading3(doc, content)
            index += 1
            continue

        if kind == "blockquote":
            while index < total and classified[index][0] == "blockquote":
                add_blockquote(doc, classified[index][1])
                index += 1
            continue

        if kind == "bullet":
            while index < total and classified[index][0] == "bullet":
                add_bullet(doc, classified[index][1])
                index += 1
            continue

        if kind == "table_row":
            table_lines: list[str] = []
            while index < total and classified[index][0] == "table_row":
                table_lines.append(classified[index][1])
                index += 1
            rows = [normalize_table_row(line) for line in table_lines]
            if len(rows) >= 2 and TABLE_DELIM_RE.match(table_lines[1]):
                rows = [rows[0], *rows[2:]]
            add_table(doc, rows)
            continue

        if kind == "hebrew_paragraph":
            if index + 1 < total and classified[index + 1][0] == "text_paragraph":
                add_contrast_pair(doc, content, classified[index + 1][1])
                index += 2
                continue
            add_hebrew_paragraph(doc, content)
            index += 1
            continue

        add_text_paragraph(doc, content)
        index += 1

    out = output or md_path.with_suffix(".docx")
    doc.save(str(out))
    return out


def parse_args(argv: Optional[Iterable[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build Hebrew lesson markdown into DOCX.")
    parser.add_argument("markdown", type=Path, nargs="+", help="Lesson cheat-sheet .md file(s)")
    parser.add_argument("-o", "--output", type=Path, help="Output DOCX path (single input only)")
    return parser.parse_args(argv)


def main(argv: Optional[Iterable[str]] = None) -> int:
    args = parse_args(argv)
    if args.output and len(args.markdown) > 1:
        print("ERROR: -o/--output only works with a single input", file=sys.stderr)
        return 1

    for md in args.markdown:
        if not md.exists():
            print(f"ERROR: not found: {md}", file=sys.stderr)
            return 1
        out = build(md, args.output)
        print(f"DOCX: {out}", flush=True, file=open(sys.stdout.fileno(), "w", encoding="utf-8", closefd=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
