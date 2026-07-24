#!/usr/bin/env python3
from __future__ import annotations

import argparse
import math
import os
import re
import subprocess
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, List, Optional, Sequence

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


HEBREW_RE = re.compile(r"[\u0590-\u05FF]")
HEBREW_FRAGMENT_RE = re.compile(r"([\u0590-\u05FF]+)")
LTR_LETTER_RE = re.compile(r"[A-Za-z\u0400-\u052F]")
PAGE_RE = re.compile(r"^#\s+Страница\s+(\d+)\s*$")
SUBHEADING_RE = re.compile(r"^##\s+")
MD_TABLE_ROW_RE = re.compile(r"^\|.*\|\s*$")
# Горизонтальная линия markdown (разделитель страниц ---). В документе не нужна:
# разрывы между страницами делаются секциями Word, а не текстом.
HR_RE = re.compile(r"^-{3,}$")
DELIM_RE = re.compile(r"^\|\s*:?[-]+:?(?:\s*\|\s*:?[-]+:?)+\s*\|\s*$")
MD_NAME_RE = re.compile(r"HP_ch(\d+)_(\d+)_(\d+)_translate\.md$", re.IGNORECASE)
IMAGE_NAME_RE = re.compile(
    r"^HP_ch(?P<chapter>[1-9]\d*)_page_(?P<page>[1-9]\d*)\.png$"
)
# Фрагмент в круглых скобках (без вложенности). В первой колонке таблицы
# «Различия» текст вне скобок — иврит, а перевод-глосса в скобках — русский.
PAREN_RE = re.compile(r"(\([^()]*\))")


@dataclass
class Block:
    kind: str  # paragraph | table
    content: object
    section: Optional[str] = None  # заголовок секции (## ...), из которой пришёл блок


@dataclass
class PageContent:
    number: int
    blocks: List[Block] = field(default_factory=list)


@dataclass
class BuildResult:
    output_docx: Path
    render_dir: Optional[Path] = None


@dataclass(frozen=True)
class InlineSpan:
    text: str
    bold: bool = False
    code: bool = False


def contains_hebrew(text: str) -> bool:
    return bool(HEBREW_RE.search(text or ""))


def parse_inline_markdown(text: str) -> List[InlineSpan]:
    """Parse the inline Markdown supported by translated chapter files.

    ``**...**`` becomes a bold run. Backticks are semantic delimiters used for
    English source fragments; they are removed while the enclosed text keeps
    the typography dictated by its script and table column.
    """
    spans: List[InlineSpan] = []
    buffer: List[str] = []
    bold = False
    code = False
    i = 0

    def flush() -> None:
        if not buffer:
            return
        value = "".join(buffer)
        buffer.clear()
        if spans and spans[-1].bold == bold and spans[-1].code == code:
            previous = spans[-1]
            spans[-1] = InlineSpan(previous.text + value, bold=bold, code=code)
        else:
            spans.append(InlineSpan(value, bold=bold, code=code))

    while i < len(text):
        if text[i] == "`":
            flush()
            code = not code
            i += 1
            continue
        if not code and text.startswith("**", i):
            flush()
            bold = not bold
            i += 2
            continue
        buffer.append(text[i])
        i += 1

    flush()
    if bold or code:
        markers = []
        if bold:
            markers.append("**")
        if code:
            markers.append("`")
        raise ValueError(
            "ERROR: invalid inline markdown; unclosed marker(s): "
            + ", ".join(markers)
        )
    return spans


def validate_inline_markdown(pages: Sequence[PageContent]) -> None:
    """Fail before document creation if supported inline markup is malformed."""
    for page in pages:
        for block in page.blocks:
            if block.kind == "paragraph":
                for chunk in str(block.content).split("\n"):
                    parse_inline_markdown(chunk)
            elif block.kind == "table":
                for row in block.content:  # type: ignore[union-attr]
                    for cell in row:
                        parse_inline_markdown(cell)


def normalize_table_row(line: str) -> List[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def parse_markdown(markdown_text: str) -> List[PageContent]:
    lines = markdown_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    pages: List[PageContent] = []
    current: Optional[PageContent] = None
    current_section: Optional[str] = None
    paragraph_buffer: List[str] = []
    i = 0

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer, current
        if current is None:
            paragraph_buffer = []
            return
        text = "\n".join(paragraph_buffer).strip()
        paragraph_buffer = []
        if text:
            current.blocks.append(Block("paragraph", text))

    while i < len(lines):
        line = lines[i]
        page_match = PAGE_RE.match(line.strip())
        if page_match:
            flush_paragraph()
            current = PageContent(number=int(page_match.group(1)))
            current_section = None
            pages.append(current)
            i += 1
            continue

        if current is None:
            i += 1
            continue

        stripped = line.strip()
        if SUBHEADING_RE.match(stripped):
            flush_paragraph()
            current_section = SUBHEADING_RE.sub("", stripped).strip()
            i += 1
            continue

        if HR_RE.match(stripped):
            # Разделитель страниц --- отбрасываем (см. HR_RE).
            flush_paragraph()
            i += 1
            continue

        if MD_TABLE_ROW_RE.match(stripped):
            flush_paragraph()
            table_lines = [stripped]
            i += 1
            while i < len(lines) and MD_TABLE_ROW_RE.match(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [normalize_table_row(x) for x in table_lines]
            if len(rows) >= 2 and DELIM_RE.match(table_lines[1]):
                rows = rows[2:]  # drop markdown header + delimiter
            if rows:
                current.blocks.append(Block("table", rows, section=current_section))
            continue

        if stripped == "":
            flush_paragraph()
            i += 1
            continue

        paragraph_buffer.append(line)
        i += 1

    flush_paragraph()
    return pages


def validate_inputs(
    md_path: Path,
    pages: Sequence[PageContent],
    image_paths: Sequence[Path],
) -> dict[int, Path]:
    md_match = MD_NAME_RE.fullmatch(md_path.name)
    if md_match is None:
        raise ValueError(
            "ERROR: markdown filename must match "
            "HP_ch{CHAPTER}_{FROM}_{TO}_translate.md"
        )

    chapter, page_from, page_to = (int(value) for value in md_match.groups())
    if page_from > page_to:
        raise ValueError("ERROR: markdown page range is reversed")

    expected_pages = list(range(page_from, page_to + 1))
    markdown_pages = [page.number for page in pages]
    if markdown_pages != expected_pages:
        raise ValueError(
            "ERROR: page count or order mismatch; "
            f"expected {expected_pages}, got {markdown_pages}"
        )

    mapping: dict[int, Path] = {}
    for image_path in image_paths:
        if not image_path.exists() or not image_path.is_file():
            raise ValueError(f"ERROR: image file not found: {image_path}")

        image_match = IMAGE_NAME_RE.fullmatch(image_path.name)
        if image_match is None:
            raise ValueError(
                "ERROR: invalid image filename "
                f"{image_path.name!r}; expected HP_ch{{CHAPTER}}_page_{{PAGE}}.png"
            )

        image_chapter = int(image_match.group("chapter"))
        page_number = int(image_match.group("page"))
        if image_chapter != chapter:
            raise ValueError(
                f"ERROR: image {image_path.name!r} belongs to chapter "
                f"{image_chapter}, expected chapter {chapter}"
            )
        if page_number in mapping:
            raise ValueError(f"ERROR: duplicate image for page {page_number}")
        mapping[page_number] = image_path

    expected_set = set(expected_pages)
    actual_set = set(mapping)
    missing = sorted(expected_set - actual_set)
    extra = sorted(actual_set - expected_set)
    if missing:
        raise ValueError(
            "ERROR: image for page "
            + ", ".join(str(page) for page in missing)
            + " not found"
        )
    if extra:
        raise ValueError(
            "ERROR: unexpected image for page "
            + ", ".join(str(page) for page in extra)
        )

    return mapping


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_no_break(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_paragraph_bidi(paragraph, enabled: bool) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1" if enabled else "0")


def set_paragraph_jc(paragraph, value: str) -> None:
    """Set w:jc directly using OOXML values ('start', 'end', 'center', 'both')."""
    p_pr = paragraph._p.get_or_add_pPr()
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), value)


def ensure_rtl_run(run, rtl: bool) -> None:
    r_pr = run._r.get_or_add_rPr()
    rtl_el = r_pr.find(qn("w:rtl"))
    if rtl_el is None:
        rtl_el = OxmlElement("w:rtl")
        r_pr.append(rtl_el)
    rtl_el.set(qn("w:val"), "1" if rtl else "0")


def set_run_font(run, *, font_name: str, font_size_pt: int, rtl: bool) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    r_pr = run._r.get_or_add_rPr()

    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), font_name)

    sz = r_pr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        r_pr.append(sz)
    sz.set(qn("w:val"), str(font_size_pt * 2))

    sz_cs = r_pr.find(qn("w:szCs"))
    if sz_cs is None:
        sz_cs = OxmlElement("w:szCs")
        r_pr.append(sz_cs)
    sz_cs.set(qn("w:val"), str(font_size_pt * 2))

    ensure_rtl_run(run, rtl)
    cs_el = r_pr.find(qn("w:cs"))
    if cs_el is None:
        cs_el = OxmlElement("w:cs")
        r_pr.append(cs_el)
    cs_el.set(qn("w:val"), "1" if rtl else "0")


def set_run_bold(run) -> None:
    """Set bold for both ordinary and complex-script text."""
    run.bold = True
    r_pr = run._r.get_or_add_rPr()
    b_cs = r_pr.find(qn("w:bCs"))
    if b_cs is None:
        b_cs = OxmlElement("w:bCs")
        r_pr.append(b_cs)
    b_cs.set(qn("w:val"), "1")


def add_formatted_run(
    paragraph,
    text: str,
    *,
    font_name: str,
    font_size_pt: int,
    rtl: bool,
    bold: bool,
) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    set_run_font(run, font_name=font_name, font_size_pt=font_size_pt, rtl=rtl)
    if bold:
        set_run_bold(run)


def style_ltr_paragraph_with_hebrew_fragments(paragraph, text: str) -> None:
    """Оформляет русский/английский текст с отдельными ивритскими фрагментами.

    В пояснениях к таблицам встречаются ивритские слова внутри русского текста.
    Направление абзаца остаётся LTR, но каждый ивритский фрагмент получает
    собственный run David/RTL — иначе Word применяет к нему русский шрифт.
    """
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_bidi(paragraph, False)

    for span in parse_inline_markdown(text):
        for part in HEBREW_FRAGMENT_RE.split(span.text):
            if not part:
                continue
            if HEBREW_FRAGMENT_RE.fullmatch(part):
                add_formatted_run(
                    paragraph,
                    part,
                    font_name="David",
                    font_size_pt=18,
                    rtl=True,
                    bold=span.bold,
                )
            else:
                add_formatted_run(
                    paragraph,
                    part,
                    font_name="Times New Roman",
                    font_size_pt=12,
                    rtl=False,
                    bold=span.bold,
                )

    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15


def style_paragraph_text(paragraph, text: str, force_hebrew: Optional[bool] = None) -> None:
    text_has_hebrew = contains_hebrew(text)
    has_hebrew = text_has_hebrew if force_hebrew is None else force_hebrew
    has_ltr_letters = bool(LTR_LETTER_RE.search(text))
    if text_has_hebrew and ((force_hebrew is False) or (force_hebrew is None and has_ltr_letters)):
        style_ltr_paragraph_with_hebrew_fragments(paragraph, text)
        return

    paragraph.text = ""
    if has_hebrew:
        set_paragraph_bidi(paragraph, True)
        set_paragraph_jc(paragraph, "start")
        for span in parse_inline_markdown(text):
            add_formatted_run(
                paragraph,
                span.text,
                font_name="David",
                font_size_pt=18,
                rtl=True,
                bold=span.bold,
            )
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_bidi(paragraph, False)
        for span in parse_inline_markdown(text):
            add_formatted_run(
                paragraph,
                span.text,
                font_name="Times New Roman",
                font_size_pt=12,
                rtl=False,
                bold=span.bold,
            )

    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15


def style_hebrew_paren_cell(paragraph, text: str) -> None:
    """Ячейка «иврит + русская глосса в скобках» (1-я колонка таблицы «Различия»).

    Базовое направление — RTL (иврит). Фрагменты в круглых скобках всегда
    оформляются как русский (Times New Roman, 12, LTR) — это переводы-глоссы.
    Текст вне скобок обычно на иврите (David, 18, RTL), но если во фрагменте
    нет ивритских букв (например, сравнение с русским словом вне скобок), он
    тоже оформляется как русский. Скобки входят в русский фрагмент.
    """
    paragraph.text = ""
    set_paragraph_bidi(paragraph, True)
    set_paragraph_jc(paragraph, "start")
    for span in parse_inline_markdown(text):
        for part in PAREN_RE.split(span.text):
            if part == "":
                continue
            in_paren = part.startswith("(") and part.endswith(")")
            if in_paren or not contains_hebrew(part):
                add_formatted_run(
                    paragraph,
                    part,
                    font_name="Times New Roman",
                    font_size_pt=12,
                    rtl=False,
                    bold=span.bold,
                )
            else:
                add_formatted_run(
                    paragraph,
                    part,
                    font_name="David",
                    font_size_pt=18,
                    rtl=True,
                    bold=span.bold,
                )

    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15


def style_page_heading(paragraph, text: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = True
    set_run_font(run, font_name="Times New Roman", font_size_pt=14, rtl=False)


def set_document_defaults(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    # Стандартные «узкие» поля Word (пресет Narrow): 1.27 см со всех сторон.
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)


def get_text_width_emu(document: Document) -> int:
    section = document.sections[-1]
    return section.page_width - section.left_margin - section.right_margin


def get_text_width_twips(document: Document) -> int:
    return int(round(get_text_width_emu(document) / 635))


def add_picture_for_page(document: Document, image_path: Path) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    width = get_text_width_emu(document)
    run = p.add_run()
    run.add_picture(str(image_path), width=width)


def build_table(
    document: Document,
    rows: List[List[str]],
    col_modes: Optional[List[Optional[str]]] = None,
) -> None:
    """Строит таблицу Word из markdown-строк.

    col_modes: если задан, оформление ячейки определяется не автоопределением по
    тексту, а принудительно по индексу колонки:
      "hebrew"        — иврит (David, 18, RTL);
      "russian"       — русский (Times New Roman, 12, LTR);
      "hebrew_paren"  — базово иврит (David/RTL), а фрагменты в круглых скобках —
                        русский (Times New Roman/LTR);
      None            — автоопределение по содержимому ячейки.
    Нужно для таблиц со смешанным текстом (иврит + русский + английский в одной
    ячейке), где автоопределение по всей ячейке ошибается.
    """
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    table_width_twips = get_text_width_twips(document)
    col_width = max(1, math.floor(table_width_twips / col_count))

    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx]
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        set_row_no_break(row)
        for col_idx in range(col_count):
            cell = row.cells[col_idx]
            set_cell_width(cell, col_width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            no_wrap = tc_pr.find(qn("w:noWrap"))
            if no_wrap is not None:
                tc_pr.remove(no_wrap)

            for p in cell.paragraphs:
                p.clear()
            text = row_data[col_idx] if col_idx < len(row_data) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.right_indent = Pt(0)
            mode = col_modes[col_idx] if (col_modes is not None and col_idx < len(col_modes)) else None
            if mode == "hebrew_paren":
                style_hebrew_paren_cell(paragraph, text)
            else:
                force_hebrew = True if mode == "hebrew" else False if mode == "russian" else None
                style_paragraph_text(paragraph, text, force_hebrew=force_hebrew)

            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side in ("top", "left", "bottom", "right"):
                el = tc_mar.find(qn(f"w:{side}"))
                if el is None:
                    el = OxmlElement(f"w:{side}")
                    tc_mar.append(el)
                el.set(qn("w:w"), "100")
                el.set(qn("w:type"), "dxa")

    document.add_paragraph().paragraph_format.space_after = Pt(6)


def derive_output_name(md_path: Path) -> str:
    match = MD_NAME_RE.search(md_path.name)
    if match:
        chapter, page_from, page_to = match.groups()
        return f"Гарри Поттер глава {chapter} страницы {page_from}-{page_to}.docx"
    return f"{md_path.stem}.docx"


def render_docx(docx_path: Path, out_dir: Path) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)
    env = os.environ.copy()
    profile_dir = out_dir / "lo_profile"
    profile_dir.mkdir(parents=True, exist_ok=True)
    env["HOME"] = str(profile_dir)
    subprocess.run(
        [
            "libreoffice",
            "--headless",
            f"-env:UserInstallation=file://{profile_dir}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(out_dir),
            str(docx_path),
        ],
        check=True,
        env=env,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )


def build_docx(
    md_path: Path,
    image_paths: Sequence[Path],
    output_path: Optional[Path],
    render: bool,
) -> BuildResult:
    pages = parse_markdown(md_path.read_text(encoding="utf-8"))
    image_map = validate_inputs(md_path, pages, image_paths)
    validate_inline_markdown(pages)

    document = Document()
    set_document_defaults(document)

    for idx, page in enumerate(pages):
        if idx > 0:
            # build_table() leaves a blank spacer paragraph after a table. If
            # that spacer is the final item on a full page, Word moves it to a
            # new page; combined with page_break_before below this created an
            # entirely blank page before the next page heading.
            trailing = next(
                (element for element in reversed(document._body._body) if element.tag != qn("w:sectPr")),
                None,
            )
            if trailing is not None and trailing.tag == qn("w:p") and not trailing.findall(f".//{qn('w:t')}"):
                trailing.getparent().remove(trailing)

        heading = document.add_paragraph()
        # Page-break-before avoids a blank page when a preceding table ends
        # exactly at a page boundary; a section break in that situation can
        # be pushed to the next page and then force the heading one page over.
        if idx > 0:
            heading.paragraph_format.page_break_before = True
        style_page_heading(heading, f"Страница {page.number}")

        add_picture_for_page(document, image_map[page.number])

        for block in page.blocks:
            if block.kind == "paragraph":
                for chunk in str(block.content).split("\n"):
                    para = document.add_paragraph()
                    style_paragraph_text(para, chunk)
            elif block.kind == "table":
                # В таблице «Различия…» ячейки смешивают иврит, русский и
                # английский, поэтому автоопределение шрифта по тексту ячейки
                # ошибается. Задаём шрифт принудительно по колонке:
                # 1-я — иврит, но глоссы в круглых скобках — русским шрифтом;
                # 2-я — целиком русский (Times New Roman).
                col_modes: Optional[List[Optional[str]]] = None
                if block.section and "Различия" in block.section:
                    col_modes = ["hebrew_paren", "russian"]
                build_table(document, block.content, col_modes=col_modes)  # type: ignore[arg-type]

    out_path = output_path or (md_path.parent / derive_output_name(md_path))
    document.save(str(out_path))

    render_dir = None
    if render:
        render_dir = out_path.with_suffix("")
        render_docx(out_path, render_dir)

    return BuildResult(output_docx=out_path, render_dir=render_dir)


def parse_args(argv: Optional[Iterable[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build Harry Potter Hebrew markdown into DOCX.")
    parser.add_argument("markdown", type=Path, help="Path to HP_ch*_translate.md")
    parser.add_argument(
        "images",
        type=Path,
        nargs="+",
        help="PNG illustrations named HP_ch{CHAPTER}_page_{PAGE}.png",
    )
    parser.add_argument("-o", "--output", type=Path, help="Output DOCX path")
    parser.add_argument("--no-render", action="store_true", help="Skip LibreOffice PDF render")
    return parser.parse_args(argv)


def main(argv: Optional[Iterable[str]] = None) -> int:
    args = parse_args(argv)
    try:
        result = build_docx(
            md_path=args.markdown,
            image_paths=args.images,
            output_path=args.output,
            render=not args.no_render,
        )
    except (OSError, ValueError) as exc:
        print(str(exc), file=sys.stderr)
        return 1
    print(f"DOCX: {result.output_docx}")
    if result.render_dir:
        print(f"RENDER_DIR: {result.render_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
