import importlib.util
import base64
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


SCRIPT = Path(__file__).parents[1] / "scripts" / "build_hp_docx.py"
SPEC = importlib.util.spec_from_file_location("build_hp_docx", SCRIPT)
assert SPEC and SPEC.loader
MODULE = importlib.util.module_from_spec(SPEC)
sys.modules[SPEC.name] = MODULE
SPEC.loader.exec_module(MODULE)


class InlineMarkdownTests(unittest.TestCase):
    def test_parser_removes_markers_and_tracks_bold(self):
        spans = MODULE.parse_inline_markdown(
            "До `Diagon Alley`: **Гарри Поттер!**"
        )
        self.assertEqual("".join(span.text for span in spans), "До Diagon Alley: Гарри Поттер!")
        self.assertEqual([span.text for span in spans if span.bold], ["Гарри Поттер!"])

    def test_parser_rejects_unclosed_markers(self):
        for text in ("это **ошибка", "это `ошибка"):
            with self.subTest(text=text):
                with self.assertRaisesRegex(ValueError, "invalid inline markdown"):
                    MODULE.parse_inline_markdown(text)

    def test_hebrew_bold_uses_complex_script_bold(self):
        document = Document()
        paragraph = document.add_paragraph()
        MODULE.style_paragraph_text(paragraph, "זה **מַצְחִיק**.")

        self.assertEqual(paragraph.text, "זה מַצְחִיק.")
        bold_runs = [run for run in paragraph.runs if run.text == "מַצְחִיק"]
        self.assertEqual(len(bold_runs), 1)
        run = bold_runs[0]
        self.assertTrue(run.bold)
        self.assertEqual(run._r.rPr.find(qn("w:bCs")).get(qn("w:val")), "1")
        self.assertEqual(run._r.rPr.find(qn("w:rtl")).get(qn("w:val")), "1")
        self.assertEqual(run._r.rPr.rFonts.get(qn("w:cs")), "David")

    def test_mixed_table_cell_preserves_script_fonts_and_bold(self):
        document = Document()
        paragraph = document.add_paragraph()
        MODULE.style_ltr_paragraph_with_hebrew_fragments(
            paragraph, "«Это не **смешно**»: זה **מַצְחִיק**"
        )

        self.assertNotIn("**", paragraph.text)
        russian = next(run for run in paragraph.runs if run.text == "смешно")
        hebrew = next(run for run in paragraph.runs if run.text == "מַצְחִיק")
        self.assertTrue(russian.bold)
        self.assertTrue(hebrew.bold)
        self.assertEqual(russian.font.name, "Times New Roman")
        self.assertEqual(hebrew._r.rPr.rFonts.get(qn("w:cs")), "David")

    def test_build_docx_converts_inline_markdown_end_to_end(self):
        png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        markdown = """# Страница 1

## Иврит

זה **טוב**.

## Подстрочный перевод

| Иврит | Перевод |
|---|---|
| **טוב** | **хорошо** |

## Различия ивритского и русского переводов

| Фрагмент | Комментарий |
|---|---|
| טוב | Английский `good` близок по смыслу. |
"""
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            md_path = root / "HP_ch1_1_1_translate.md"
            image_path = root / "HP_ch1_page_1.png"
            output_path = root / "result.docx"
            md_path.write_text(markdown, encoding="utf-8")
            image_path.write_bytes(png)

            MODULE.build_docx(md_path, [image_path], output_path, render=False)

            document = Document(output_path)
            all_text = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    all_text.extend(cell.text for cell in row.cells)
            joined = "\n".join(all_text)
            self.assertNotIn("**", joined)
            self.assertNotIn("`", joined)
            self.assertIn("good", joined)
            self.assertTrue(
                any(run.text == "טוב" and run.bold for paragraph in document.paragraphs for run in paragraph.runs)
            )

    def test_invalid_inline_markdown_does_not_create_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            md_path = root / "HP_ch1_1_1_translate.md"
            image_path = root / "HP_ch1_page_1.png"
            output_path = root / "result.docx"
            md_path.write_text("# Страница 1\n\n## Иврит\n\nזה **שגוי\n", encoding="utf-8")
            image_path.write_bytes(b"placeholder")

            with self.assertRaisesRegex(ValueError, "invalid inline markdown"):
                MODULE.build_docx(md_path, [image_path], output_path, render=False)
            self.assertFalse(output_path.exists())


if __name__ == "__main__":
    unittest.main()
